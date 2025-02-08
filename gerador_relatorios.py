from docx import Document
from openpyxl import Workbook
from gerenciador_dados import listar_alunos
from rich.console import Console
from time import sleep

console = Console()

def gerar_rel(lista, op):
    if op == 1: #individual
        try:
            listar_alunos(lista)
            matricula = input("Matricula: ")
            for aluno in lista:
                if matricula == aluno['matricula'] and aluno['situacao-matricula'] == 'ativa':
                    media = sum(aluno['notas'])/3
                    situacao = 'Reprovado'
                    if media >= 7 and aluno['total-faltas'] < 50:
                        situacao = 'Aprovado'
                    doc = Document()
                    doc.add_paragraph(f"Nome: {aluno['nome']}")
                    doc.add_paragraph(f"Matrícula: {matricula}")
                    doc.add_paragraph(f"Data de Nascimento: {aluno['data-nasc']}")
                    doc.add_paragraph(f"Média: {round(media, 2)}")
                    doc.add_paragraph(f"{situacao}")
                    doc.save(f"relatorio-{aluno['nome']}.docx")
                    return True
            return False
        except:
            console.print("O arquivo está aberto, feche ele para conseguir gerar o arquivo!")
    elif op == 2: #coletivo
        try:
            doc = Document()
            doc.add_heading("Relatório Coletivo", level=1)
            for aluno in lista:
                if aluno['situacao-matricula'] == 'ativa':
                    media = sum(aluno['notas'])/3
                    situacao = 'Reprovado'
                    if media >= 7 and aluno['total-faltas'] < 50:
                        situacao = 'Aprovado'
                    doc.add_heading(f"Nome: {aluno['nome']}", level=2)
                    doc.add_paragraph(f"Matrícula: {aluno['matricula']}")
                    doc.add_paragraph(f"Data de Nascimento: {aluno['data-nasc']}")
                    doc.add_paragraph(f"Média: {round(media, 2)}")
                    doc.add_paragraph(f"{situacao}\n")
            doc.save("relatorio-geral.docx")
            return True
        except:
            console.print("O arquivo está aberto, feche ele para conseguir gerar o arquivo!")
            return False


def gerar_planilha(lista):
    wb = Workbook()
    ws = wb.active
    ws.title = "Planilha Alunos"
    ws.append(['Nome', 'Matricula', 'Data de nascimento', 
            'Nota 1', 'Nota 2', 'Nota 3', 'Média', 'Total de faltas', 
            'Status de aprovação'])
    for aluno in lista:
        if aluno['situacao-matricula'] == 'ativa':
            media = sum(aluno['notas'])/3
            situacao = 'reprovado'
            if media >= 7 and aluno['total-faltas'] < 50:
                situacao = 'aprovado'
            ws.append([aluno['nome'], aluno['matricula'], 
                    aluno['data-nasc'], aluno['notas'][0], 
                    aluno['notas'][1], aluno['notas'][2],
                    round(media, 2), aluno['total-faltas'],
                    situacao])
    wb.save('planilha-alunos.xlsx')
    return True
